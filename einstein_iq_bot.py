import db
import config
import random
import aiofiles
import docx
import os
from aiogram import Bot
from aiogram import Dispatcher
from aiogram import executor
from aiogram.types import InlineKeyboardButton
from aiogram.types import InlineKeyboardMarkup
from aiogram.types import CallbackQuery
from aiogram.types import Message
from json import dumps
from json import loads
from aiogram import Bot, Dispatcher
from aiogram.utils import executor


def reset(uid: int):
    db.set_in_process(uid, False)
    db.set_selected_ticket(uid, 0)
    db.change_questions_passed(uid, 0)
    db.change_questions_message(uid, 0)
    db.change_current_question(uid, 0)


bot = Bot(token=config.TOKEN)
dp = Dispatcher(bot=bot)


@dp.message_handler(commands=["play"])
async def go_handler(message: Message):
    if not db.is_exists(message.from_user.id):
        db.add(message.from_user.id)
    if db.is_in_process(message.from_user.id):
        await bot.send_message(
            message.from_user.id,
            "🚫*Вы не можете начать тест, потому что вы его уже проходите*\\.",
            parse_mode="MarkdownV2"
        )
        return


    document = docx.Document()


    ticket_id = db.get_selected_ticket(message.from_user.id)
    if ticket_id:
        selected_data = f"ticket/bilet{ticket_id}.json"
    else:
        ticket_id = random.randint(1, 5)
        selected_data = f"ticket/bilet{ticket_id}.json"
        db.set_selected_ticket(message.from_user.id, ticket_id)
    async with aiofiles.open(selected_data, "r", encoding="utf-8") as file:
        questions = await file.read()
        questions = loads(questions)


    def compose_markup(question: int):
        km = InlineKeyboardMarkup(row_width=3)
        for i in range(len(questions[question]["variants"])):
            cd = {"question": question, "answer": i}
            km.insert(
                InlineKeyboardButton(
                    questions[question]["variants"][i],
                    callback_data=dumps(cd)
                )
            )
        return km


    @dp.callback_query_handler(lambda c: True)
    async def answer_handler(callback: CallbackQuery):
        ticket_id = db.get_selected_ticket(message.from_user.id)
        if ticket_id:
            selected_data = f"ticket/bilet{ticket_id}.json"
        else:
            ticket_id = random.randint(1, 5)
            selected_data = f"ticket/bilet{ticket_id}.json"
            db.set_selected_ticket(message.from_user.id, ticket_id)
        async with aiofiles.open(selected_data, "r", encoding="utf-8") as file:
            questions = await file.read()
            questions = loads(questions)
        data = loads(callback.data)
        q = data["question"]
        is_correct = questions[q]["correct_answer"] - 1 == data["answer"]
        namedocument = (f"{callback.from_user.id}.docx")
        document.add_paragraph(questions[q]["text"])
        if is_correct==True:
            p = document.add_paragraph()
            p.add_run('Ответ пользователя: ').bold=True
            p.add_run((str(data["answer"]++1))).bold=True
            p.add_run(' - верно').bold=True
            document.add_paragraph('\n')
        else: 
            p = document.add_paragraph()
            p.add_run('Ответ пользователя: ').bold=True
            p.add_run((str(data["answer"]++1))).bold=True
            p.add_run(' - не верно').bold=True
            document.add_paragraph('\n')
        document.save(namedocument)
        passed = db.get_questions_passed(callback.from_user.id)
        msg = db.get_questions_message(callback.from_user.id)
        if is_correct:
            passed += 1
            db.change_questions_passed(callback.from_user.id, passed)
        if q + 1 > len(questions) - 1:
            if passed <= 7:
                reset(callback.from_user.id)
                await bot.delete_message(callback.from_user.id, msg)
                await bot.send_message(
                    callback.from_user.id,
                    f"🚫 Вы не прошли это тест\\!\n\n✅ Правильных ответов\\: *{passed} из {len(questions)}*\\.\n\n🔄 *Пройти тест снова* \\- /play",
                    parse_mode="MarkdownV2"
                )
                doc = docx.Document(namedocument)
                for para in doc.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace('\\', '').replace('*', '')                  
                doc.save(namedocument)
                await bot.send_document(callback.from_user.id,document=open(namedocument,'rb'))
                os.remove(namedocument)
            if passed >= 8:
                reset(callback.from_user.id)
                await bot.delete_message(callback.from_user.id, msg)
                await bot.send_message(
                    callback.from_user.id,
                    f"🎉 *Ура*, вы прошли это тест\\!\n\n✅ Правильных ответов\\: *{passed} из {len(questions)}*\\.",
                    parse_mode="MarkdownV2"
                )
                doc = docx.Document(namedocument)
                for para in doc.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace('\\', '').replace('*', '')
                doc.save(namedocument)
                await bot.send_document(callback.from_user.id,document=open(namedocument,'rb'))
                os.remove(namedocument)
            return


        await bot.edit_message_text(
            questions[q + 1]["text"],
            callback.from_user.id,
            msg,
            reply_markup=compose_markup(q + 1),
            parse_mode="MarkdownV2"
        )
    

    db.set_in_process(message.from_user.id, True)
    msg = await bot.send_message(
        message.from_user.id,
        questions[0]["text"],
        reply_markup=compose_markup(0),
        parse_mode="MarkdownV2"
    )

    db.change_questions_message(message.from_user.id, msg.message_id)
    db.change_current_question(message.from_user.id, 0)
    db.change_questions_passed(message.from_user.id, 0)


@dp.message_handler(commands=["finish"])
async def quit_handler(message: Message):
    if not db.is_in_process(message.from_user.id):
        await bot.send_message(message.from_user.id, "❗️Вы ещё *не начали тест*\\.", parse_mode="MarkdownV2")
        return
    reset(message.from_user.id)
    await bot.send_message(message.from_user.id, "❗️Вы *остановили тест*\\.", parse_mode="MarkdownV2")


@dp.message_handler(commands=["start"])
async def start(message: Message):
    await message.answer("👋 *Привет\\!* \n🧠 *Это внеочередная проверка знаний по новым правилам технической эксплуатации электроустановок потребителей электрической энергии\\.*\n\n📝 Нужно будет ответить на *10 вопросов*\\. \n⏱ Если будете думать как следует, тест займет *около 10 минут*\\. \n\n⁉️ *8* — правильных ответов для успешной сдачи тестирования\\. \n📄 К каждому вопросу предлагаются *несколько вариантов ответов*\\. \n\n🔍 Верным является * только один * вариант, его вам и надо выбрать\\.\n\n*Начать тест* \\- /play\n*Закончить тест* \\- /finish\n*Техническая поддержка* \\- /help", parse_mode="MarkdownV2")


@dp.message_handler(commands=['help'])
async def cmd_answer(message: Message):
    await message.answer("⁉️<b> Если у вас есть проблемы.</b> \n✉️ <b>Напишите мне</b> <a href='https://t.me/mistrcron'>@mistrcron</a><b>.</b>", disable_web_page_preview=True, parse_mode="HTML")


def main() -> None:
    executor.start_polling(dp, skip_updates=True)


if __name__ == "__main__":
    main()